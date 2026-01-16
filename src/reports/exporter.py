"""
报告导出器
支持将 LegalDocumentReport 导出为多种格式
"""

import os
from abc import ABC, abstractmethod
from typing import Optional, Literal
from pathlib import Path

from ..models.report_schema import LegalDocumentReport


class BaseExporter(ABC):
    """
    报告导出器基类

    所有导出器都应继承此类并实现 export 方法
    """

    def __init__(self, output_dir: str = "reports"):
        """
        初始化导出器

        Args:
            output_dir: 输出目录
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    @abstractmethod
    def export(
        self,
        report: LegalDocumentReport,
        filename: Optional[str] = None
    ) -> str:
        """
        导出报告

        Args:
            report: 报告对象
            filename: 文件名（不含扩展名）

        Returns:
            导出文件的完整路径
        """
        pass

    def _generate_filename(
        self,
        report: LegalDocumentReport,
        filename: Optional[str],
        extension: str
    ) -> str:
        """生成文件名"""
        if filename:
            base_name = filename
        else:
            base_name = f"{report.document_name}_{report.analysis_timestamp.strftime('%Y%m%d_%H%M%S')}"

        # 确保文件名安全
        base_name = self._sanitize_filename(base_name)

        return str(self.output_dir / f"{base_name}.{extension}")

    def _sanitize_filename(self, filename: str) -> str:
        """清理文件名中的非法字符"""
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        return filename


class JSONExporter(BaseExporter):
    """JSON 导出器"""

    def export(
        self,
        report: LegalDocumentReport,
        filename: Optional[str] = None
    ) -> str:
        """
        导出为 JSON

        Args:
            report: 报告对象
            filename: 文件名

        Returns:
            JSON 文件路径
        """
        filepath = self._generate_filename(report, filename, "json")

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(report.to_json(indent=2, ensure_ascii=False))

        return filepath


class MarkdownExporter(BaseExporter):
    """Markdown 导出器"""

    def export(
        self,
        report: LegalDocumentReport,
        filename: Optional[str] = None
    ) -> str:
        """
        导出为 Markdown

        Args:
            report: 报告对象
            filename: 文件名

        Returns:
            Markdown 文件路径
        """
        filepath = self._generate_filename(report, filename, "md")

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(report.get_full_markdown())

        return filepath


class PDFExporter(BaseExporter):
    """
    PDF 导出器

    使用 WeasyPrint 将 Markdown/HTML 转换为 PDF
    """

    def export(
        self,
        report: LegalDocumentReport,
        filename: Optional[str] = None
    ) -> str:
        """
        导出为 PDF

        Args:
            report: 报告对象
            filename: 文件名

        Returns:
            PDF 文件路径
        """
        filepath = self._generate_filename(report, filename, "pdf")

        try:
            # 尝试使用 WeasyPrint
            import weasyprint
            from io import StringIO

            # 生成 HTML
            html_content = self._generate_html(report)

            # 转换为 PDF
            weasyprint.HTML(string=html_content).write_pdf(
                target=filepath,
                stylesheets=[self._get_css()]
            )

            return filepath

        except ImportError:
            # 降级：保存为 Markdown 并提示
            print("⚠️ WeasyPrint 未安装，降级为 Markdown 格式")
            md_exporter = MarkdownExporter(self.output_dir)
            return md_exporter.export(report, filename)

    def _generate_html(self, report: LegalDocumentReport) -> str:
        """生成 HTML 内容"""
        md_content = report.get_full_markdown()

        # 简单的 Markdown 到 HTML 转换
        # 实际生产中建议使用 markdown2 或 mistune
        html_content = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{report.document_name} - 法律分析报告</title>
</head>
<body>
    <div class="container">
        <pre style="white-space: pre-wrap; font-family: 'Microsoft YaHei', sans-serif;">
{md_content}
        </pre>
    </div>
</body>
</html>
"""
        return html_content

    def _get_css(self) -> str:
        """获取 CSS 样式"""
        return """
        @page {
            size: A4;
            margin: 2cm;
        }

        body {
            font-family: 'Microsoft YaHei', 'SimSun', sans-serif;
            font-size: 12px;
            line-height: 1.6;
        }

        h1 {
            font-size: 24px;
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }

        h2 {
            font-size: 18px;
            color: #34495e;
            margin-top: 20px;
        }

        h3 {
            font-size: 14px;
            color: #7f8c8d;
        }

        .critical {
            color: #e74c3c;
            font-weight: bold;
        }

        .high {
            color: #e67e22;
        }

        .medium {
            color: #f39c12;
        }

        .low {
            color: #27ae60;
        }
        """


class DOCXExporter(BaseExporter):
    """
    DOCX 导出器

    使用 python-docx 生成 Word 文档
    """

    def export(
        self,
        report: LegalDocumentReport,
        filename: Optional[str] = None
    ) -> str:
        """
        导出为 DOCX

        Args:
            report: 报告对象
            filename: 文件名

        Returns:
            DOCX 文件路径
        """
        filepath = self._generate_filename(report, filename, "docx")

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            # 创建文档
            doc = Document()

            # 标题
            title = doc.add_heading(f"{report.document_name} - 法律分析报告", 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 元信息
            self._add_metadata(doc, report)

            # 执行摘要
            self._add_executive_summary(doc, report)

            # 详细分析
            self._add_detailed_analysis(doc, report)

            # 证据来源
            self._add_evidence_sources(doc, report)

            # 保存
            doc.save(filepath)

            return filepath

        except ImportError:
            print("⚠️ python-docx 未安装，降级为 Markdown 格式")
            md_exporter = MarkdownExporter(self.output_dir)
            return md_exporter.export(report, filename)

    def _add_metadata(self, doc, report: LegalDocumentReport):
        """添加元信息"""
        doc.add_heading("报告信息", 1)

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        metadata = [
            ("文档名称", report.document_name),
            ("生成时间", report.analysis_timestamp.strftime('%Y-%m-%d %H:%M:%S')),
            ("分析版本", report.analysis_version),
            ("分析耗时", f"{report.analysis_duration_seconds:.2f} 秒"),
            ("综合评分", f"{report.executive_summary.overall_rating}/10")
        ]

        for i, (key, value) in enumerate(metadata):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(value)

    def _add_executive_summary(self, doc, report: LegalDocumentReport):
        """添加执行摘要"""
        doc.add_heading("执行摘要", 1)

        summary = report.executive_summary

        # 评分
        p = doc.add_paragraph()
        p.add_run("综合评分: ").bold = True
        p.add_run(f"{summary.overall_rating}/10")
        p.add_run(f"\n{summary.rating_explanation}")

        # 关键风险
        doc.add_heading("关键风险", 2)
        for risk in summary.key_risks:
            doc.add_paragraph(risk, style='List Bullet')

        # 快速建议
        doc.add_heading("快速建议", 2)
        for rec in summary.quick_recommendations:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"[{rec.priority.value.upper()}] ").bold = True
            p.add_run(f"{rec.action_item} - {rec.urgency}")

    def _add_detailed_analysis(self, doc, report: LegalDocumentReport):
        """添加详细分析"""
        doc.add_page_break()
        doc.add_heading("详细分析", 1)

        analysis = report.detailed_analysis

        # 概览
        doc.add_paragraph(
            f"总条款数: {analysis.total_clauses} | "
            f"有问题条款: {analysis.clauses_with_issues} | "
            f"合规率: {analysis.compliance_rate:.1f}%"
        )

        # 条款分析
        doc.add_heading("条款分析", 2)
        for clause in analysis.clause_breakdown:
            doc.add_heading(f"条款 {clause.clause_id}", 3)

            if clause.clause_title:
                doc.add_paragraph(clause.clause_title).bold = True

            doc.add_paragraph(f"类型: {clause.clause_type}")
            doc.add_paragraph(f"风险等级: {clause.risk_level.value}")
            doc.add_paragraph(f"风险分数: {clause.risk_score}/100")

            if clause.issues_identified:
                doc.add_paragraph("问题:")
                for issue in clause.issues_identified:
                    doc.add_paragraph(issue, style='List Bullet')

            if clause.suggestions:
                doc.add_paragraph("建议:")
                for suggestion in clause.suggestions:
                    doc.add_paragraph(suggestion, style='List Bullet')

    def _add_evidence_sources(self, doc, report: LegalDocumentReport):
        """添加证据来源"""
        doc.add_page_break()
        doc.add_heading("证据来源", 1)

        for source in report.evidence_sources:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"[{source.source_type.value}] ").bold = True
            p.add_run(source.content)
            p.add_run(f" (相关性: {source.relevance_score:.2f})").italic = True


class ReportExporter:
    """
    报告导出器

    统一接口，支持多种格式导出
    """

    def __init__(self, output_dir: str = "reports"):
        """
        初始化导出器

        Args:
            output_dir: 输出目录
        """
        self.output_dir = output_dir

        # 导出器实例
        self.json_exporter = JSONExporter(output_dir)
        self.md_exporter = MarkdownExporter(output_dir)
        self.pdf_exporter = PDFExporter(output_dir)
        self.docx_exporter = DOCXExporter(output_dir)

    def export(
        self,
        report: LegalDocumentReport,
        formats: list = None,
        filename: Optional[str] = None
    ) -> dict:
        """
        导出报告为多种格式

        Args:
            report: 报告对象
            formats: 格式列表 ["json", "pdf", "docx", "md"]
            filename: 文件名（不含扩展名）

        Returns:
            导出结果字典 {format: filepath}
        """
        if formats is None:
            formats = ["json", "md"]

        results = {}

        for fmt in formats:
            try:
                if fmt == "json":
                    filepath = self.json_exporter.export(report, filename)
                elif fmt == "md":
                    filepath = self.md_exporter.export(report, filename)
                elif fmt == "pdf":
                    filepath = self.pdf_exporter.export(report, filename)
                elif fmt == "docx":
                    filepath = self.docx_exporter.export(report, filename)
                else:
                    print(f"⚠️ 不支持的格式: {fmt}")
                    continue

                results[fmt] = filepath
                print(f"✅ {fmt.upper()} 导出成功: {filepath}")

            except Exception as e:
                print(f"⚠️ {fmt.upper()} 导出失败: {e}")
                results[fmt] = None

        return results

    def export_all(
        self,
        report: LegalDocumentReport,
        filename: Optional[str] = None
    ) -> dict:
        """
        导出所有支持的格式

        Args:
            report: 报告对象
            filename: 文件名

        Returns:
            导出结果字典
        """
        return self.export(
            report=report,
            formats=["json", "md", "pdf", "docx"],
            filename=filename
        )


# 工厂函数
def create_report_exporter(output_dir: str = "reports") -> ReportExporter:
    """
    创建报告导出器

    Args:
        output_dir: 输出目录

    Returns:
        ReportExporter 实例
    """
    return ReportExporter(output_dir=output_dir)
